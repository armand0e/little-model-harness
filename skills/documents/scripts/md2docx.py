"""Convert Markdown to a professionally styled .docx.

Usage: python md2docx.py input.md output.docx

Optional front-matter lines (before any other content):
  Title: / Subtitle: / Author: / Date: / Accent: #1F4E79
Title: produces a cover page; Accent: recolors headings and rules.
"""
import datetime
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BOLD_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")
XML_INVALID_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F]")
DEFAULT_ACCENT = "1F4E79"


def clean_xml(text):
    return XML_INVALID_RE.sub("", text)


def add_runs(paragraph, text):
    """Render **bold**, *italic* and `code` inline."""
    text = clean_xml(text)
    for part in BOLD_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def is_table_sep(line):
    return bool(re.match(r"^\s*\|?[\s:|-]+\|[\s:|-]*$", line)) and "-" in line


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def shade(element, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    element.append(shd)


def shade_cell(cell, hex_color):
    shade(cell._tc.get_or_add_tcPr(), hex_color)


def shade_paragraph(paragraph, hex_color):
    shade(paragraph._p.get_or_add_pPr(), hex_color)


def bottom_rule(paragraph, hex_color, size="14"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def left_rule(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "10")
    left.set(qn("w:color"), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)


def add_page_number_footer(doc, accent):
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    for field in ("begin", None, "end"):
        element = OxmlElement("w:fldChar" if field else "w:instrText")
        if field:
            element.set(qn("w:fldCharType"), field)
        else:
            element.set(qn("xml:space"), "preserve")
            element.text = " PAGE "
        run._r.append(element)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("808080")


def apply_styles(doc, accent):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.12
    specs = {"Heading 1": (16, accent, 16, 6),
             "Heading 2": (13, accent, 12, 4),
             "Heading 3": (11.5, "404040", 10, 3),
             "Heading 4": (11, "404040", 8, 2)}
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(1)
        section.top_margin = section.bottom_margin = Inches(0.9)


def build_cover(doc, meta, accent):
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    run = title.add_run(meta["title"])
    run.font.name = "Georgia"
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("262626")
    bottom_rule(title, accent, size="20")
    title.paragraph_format.space_after = Pt(14)
    if meta.get("subtitle"):
        subtitle = doc.add_paragraph()
        run = subtitle.add_run(meta["subtitle"])
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor.from_string("595959")
    line_parts = [part for part in (
        meta.get("author"),
        meta.get("date") or datetime.date.today().strftime("%B %d, %Y"),
    ) if part]
    if line_parts:
        info = doc.add_paragraph()
        run = info.add_run("  ·  ".join(line_parts))
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string("808080")
        info.paragraph_format.space_before = Pt(28)
    doc.add_page_break()


def parse_front_matter(lines):
    meta = {}
    index = 0
    keys = {"title", "subtitle", "author", "date", "accent"}
    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped:
            index += 1
            continue
        m = re.match(r"^(\w+):\s*(.+)$", stripped)
        if m and m.group(1).lower() in keys:
            meta[m.group(1).lower()] = clean_xml(m.group(2).strip())
            index += 1
            continue
        break
    return meta, lines[index:]


def convert(md_path, docx_path):
    raw_lines = Path(md_path).read_text(encoding="utf-8").splitlines()
    meta, lines = parse_front_matter(raw_lines)
    accent = re.sub(r"[^0-9A-Fa-f]", "", meta.get("accent", ""))[:6].upper() \
        or DEFAULT_ACCENT

    doc = Document()
    apply_styles(doc, accent)
    add_page_number_footer(doc, accent)
    if meta.get("title"):
        build_cover(doc, meta, accent)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # fenced code block: shaded mono paragraph
        if stripped.startswith("```"):
            i += 1
            code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            run = p.add_run(clean_xml("\n".join(code)))
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            shade_paragraph(p, "F2F1EE")
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(10)
            continue

        # table: accent header, banded rows
        if stripped.startswith("|") and i + 1 < len(lines) \
                and is_table_sep(lines[i + 1]):
            headers = split_row(stripped)
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            for j, header in enumerate(headers):
                cell = table.rows[0].cells[j]
                shade_cell(cell, accent)
                cell.text = ""
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run(header.strip("*"))
                run.bold = True
                run.font.color.rgb = RGBColor.from_string("FFFFFF")
            for row_index, row in enumerate(rows):
                cells = table.add_row().cells
                for j, value in enumerate(row[:len(headers)]):
                    if row_index % 2 == 1:
                        shade_cell(cells[j], "F3F2EF")
                    cells[j].text = ""
                    add_runs(cells[j].paragraphs[0], value)
            doc.add_paragraph()
            continue

        if stripped == "\\newpage":
            doc.add_page_break()
            i += 1
            continue

        if stripped in ("---", "***", "___"):
            bottom_rule(doc.add_paragraph(), "BFBDB8", size="6")
            i += 1
            continue

        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            doc.add_heading(clean_xml(m.group(2).strip()),
                            level=len(m.group(1)))
            i += 1
            continue

        # block quote: accent left rule, muted italic (merge > lines)
        if stripped.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip().lstrip("> ").strip())
                i += 1
            p = doc.add_paragraph()
            left_rule(p, accent)
            p.paragraph_format.left_indent = Inches(0.15)
            add_runs(p, " ".join(part for part in quote_lines if part))
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor.from_string("595959")
            continue

        m = re.match(r"^(\s*)([-*]|\d+[.)])\s+(.*)$", line)
        if m:
            indent = len(m.group(1)) // 2
            numbered = m.group(2)[0].isdigit()
            base = "List Number" if numbered else "List Bullet"
            style_name = base + ("" if indent == 0 else f" {min(indent + 1, 3)}")
            try:
                p = doc.add_paragraph(style=style_name)
            except KeyError:
                p = doc.add_paragraph(style=base)
            add_runs(p, m.group(3))
            i += 1
            continue

        para = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(
                r"^(\s*([-*]|\d+[.)])\s|#{1,4}\s|\||>|```)", lines[i]):
            para.append(lines[i].strip())
            i += 1
        p = doc.add_paragraph()
        add_runs(p, " ".join(para))

    output = Path(docx_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"Saved {docx_path}")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        sys.exit("Usage: python md2docx.py input.md output.docx")
    convert(sys.argv[1], sys.argv[2])
