from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pptx import Presentation
from pypdf import PdfWriter


ROOT = Path(__file__).parents[1]


def run_script(relative: str, *args: object) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [sys.executable, str(ROOT / relative), *(str(arg) for arg in args)],
        capture_output=True, text=True, encoding="utf-8", errors="replace",
        timeout=20,
    )


def test_document_helpers_create_and_read_docx(tmp_path: Path):
    source = tmp_path / "report.md"
    output = tmp_path / "nested" / "report.docx"
    source.write_text(
        "# Re\x08port\n\nA **bold** sum\x00mary.\n\n| A | B |\n|---|---|\n| 1 | 2 |",
        encoding="utf-8",
    )
    made = run_script("skills/documents/scripts/md2docx.py", source, output)
    assert made.returncode == 0, made.stderr
    doc = Document(output)
    assert any(p.text == "Report" for p in doc.paragraphs)
    assert len(doc.tables) == 1

    read = run_script("skills/documents/scripts/read_docx.py", output)
    assert read.returncode == 0
    assert "# Report" in read.stdout
    assert "A | B" in read.stdout


def test_pdf_reader_rejects_invalid_page_ranges(tmp_path: Path):
    path = tmp_path / "blank.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    with path.open("wb") as f:
        writer.write(f)

    invalid = run_script("skills/documents/scripts/read_pdf.py", path, "0")
    assert invalid.returncode != 0
    assert "outside this 1-page PDF" in invalid.stderr

    valid = run_script("skills/documents/scripts/read_pdf.py", path, "1")
    assert valid.returncode == 0
    assert "page 1 of 1" in valid.stdout


def test_spreadsheet_helpers_validate_and_deduplicate_sheet_names(
        tmp_path: Path):
    spec = tmp_path / "book.json"
    output = tmp_path / "nested" / "book.xlsx"
    spec.write_text(json.dumps({"sheets": [
        {"name": "Bad/Name", "headers": ["Value"], "rows": [["safe\x08value"]],
         "widths": {"A": 500}},
        {"name": "Bad/Name", "headers": ["Formula"], "rows": [["=1+1"]]},
    ]}), encoding="utf-8")
    made = run_script("skills/spreadsheets/scripts/make_xlsx.py", spec, output)
    assert made.returncode == 0, made.stderr
    wb = load_workbook(output, data_only=False)
    assert wb.sheetnames == ["Bad_Name", "Bad_Name (2)"]
    assert wb["Bad_Name"].column_dimensions["A"].width == 255
    assert wb["Bad_Name"]["A2"].value == "safevalue"
    assert wb["Bad_Name (2)"]["A2"].value == "=1+1"

    read = run_script("skills/spreadsheets/scripts/read_xlsx.py", output)
    assert read.returncode == 0
    assert "Sheet: Bad_Name" in read.stdout

    spec.write_text('{"sheets": [{"rows": [[{"bad": true}]]}]}',
                    encoding="utf-8")
    rejected = run_script("skills/spreadsheets/scripts/make_xlsx.py", spec, output)
    assert rejected.returncode != 0
    assert "cells must be scalar" in rejected.stderr


def test_presentation_helpers_create_read_and_report_slide_count(
        tmp_path: Path):
    source = tmp_path / "deck.md"
    output = tmp_path / "nested" / "deck.pptx"
    source.write_text(
        "# De\x08mo\nSubtitle\n\n## First\n- One\n- Two\nNotes: speaker\x00 note",
        encoding="utf-8",
    )
    made = run_script("skills/presentations/scripts/md2pptx.py", source, output)
    assert made.returncode == 0, made.stderr
    assert "(2 slides)" in made.stdout
    prs = Presentation(output)
    assert len(prs.slides) == 2

    read = run_script("skills/presentations/scripts/read_pptx.py", output)
    assert read.returncode == 0
    assert "Slide 2" in read.stdout
    assert "speaker note" in read.stdout


def test_computer_helper_validates_non_destructive_commands():
    missing = run_script("skills/computer/scripts/computer.py", "press")
    assert missing.returncode != 0
    assert "requires 1 argument" in missing.stderr

    wait = run_script("skills/computer/scripts/computer.py", "wait", -1)
    assert wait.returncode == 0
    assert "Waited 0s" in wait.stdout
